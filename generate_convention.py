"""
Generate the ADCM–IPEST partnership convention as DOCX and PDF.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_DOCX = r"c:\Users\JWHR\Desktop\coste\Convention_ADCM_IPEST.docx"
OUTPUT_PDF  = r"c:\Users\JWHR\Desktop\coste\Convention_ADCM_IPEST.pdf"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_font(run, bold=False, italic=False, size=12, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6, line_spacing=1.15):
    fmt = para.paragraph_format
    fmt.alignment = alignment
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing


def add_paragraph(doc, text="", bold=False, italic=False, size=12,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    para_format(p, alignment=alignment, space_before=space_before,
                space_after=space_after)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size)
    return p


def add_heading(doc, text, level_size=14, space_before=12, space_after=6):
    p = doc.add_paragraph()
    para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=space_before, space_after=space_after)
    run = p.add_run(text)
    set_font(run, bold=True, size=level_size)
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    para_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=4)
    # Replace bullet symbol with em dash
    run = p.add_run(text)
    set_font(run, size=size)
    # Override list style to use em dash
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.runs[0].text = ""
    run2 = p.add_run("— " + text)
    set_font(run2, size=size)
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_doc_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def center_run(doc, parts, sizes=None, bolds=None, space_before=0, space_after=6):
    """parts is list of strings, sizes list of sizes (same length)."""
    p = doc.add_paragraph()
    para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=space_before, space_after=space_after)
    for i, text in enumerate(parts):
        run = p.add_run(text)
        sz = sizes[i] if sizes else 12
        bd = bolds[i] if bolds else False
        set_font(run, bold=bd, size=sz)
    return p


def add_separator(doc):
    p = doc.add_paragraph()
    para_format(p, space_before=4, space_after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_docx():
    doc = Document()
    set_doc_margins(doc)

    # Remove default styles spacing
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ===========================================================
    # TITLE PAGE
    # ===========================================================
    for _ in range(3):
        add_paragraph(doc, space_after=4)

    center_run(doc, ["CONVENTION DE PARTENARIAT"],
               sizes=[18], bolds=[True], space_before=0, space_after=10)
    add_separator(doc)
    add_paragraph(doc, space_after=8)

    center_run(doc, ["entre"], sizes=[12], bolds=[False], space_after=6)
    add_paragraph(doc, space_after=4)

    center_run(doc,
               ["L'Association du Développement Culturel à La Marsa\n« ADCM »"],
               sizes=[13], bolds=[True], space_after=6)

    center_run(doc, ["et"], sizes=[12], bolds=[False], space_after=6)

    center_run(doc,
               ["L'Institut Préparatoire aux Études Scientifiques et Techniques\n« IPEST »"],
               sizes=[13], bolds=[True], space_after=14)

    add_separator(doc)
    add_paragraph(doc, space_after=8)

    center_run(doc, ["La Marsa — 2026"], sizes=[12], bolds=[False], space_after=4)

    add_page_break(doc)

    # ===========================================================
    # ENTRE LES SOUSSIGNÉS
    # ===========================================================
    add_heading(doc, "ENTRE LES SOUSSIGNÉS", level_size=14, space_before=0, space_after=10)

    add_paragraph(doc,
        "L'Association du Développement Culturel à La Marsa « ADCM », représentée par "
        "Madame Soumaya KHAMMAR, Présidente de l'Association, ayant tous pouvoirs aux fins "
        "de signature des présentes,",
        space_after=4)

    add_paragraph(doc, "ci-après désignée « l'Association » ou « ADCM »,",
                  italic=True, space_after=4)
    add_paragraph(doc, "d'une part,", space_after=10)

    center_run(doc, ["et"], sizes=[12], bolds=[False], space_after=10)

    add_paragraph(doc,
        "L'Institut Préparatoire aux Études Scientifiques et Techniques « IPEST », "
        "représenté par Monsieur/Madame ………………………………………………………, "
        "Directeur(trice) de l'Institut, ayant tous pouvoirs aux fins de signature des présentes,",
        space_after=4)

    add_paragraph(doc, "ci-après désigné « l'Institut » ou « IPEST »,",
                  italic=True, space_after=4)
    add_paragraph(doc, "d'autre part,", space_after=10)

    add_paragraph(doc,
        "Les soussignés étant ci-après désignés conjointement « les Parties ».",
        italic=True, space_after=14)

    add_separator(doc)

    # ===========================================================
    # PRÉAMBULE
    # ===========================================================
    add_heading(doc, "PRÉAMBULE", level_size=14, space_before=12, space_after=10)

    add_heading(doc, "Présentation de l'Association ADCM", level_size=12,
                space_before=8, space_after=6)

    add_paragraph(doc,
        "L'Association du Développement Culturel à La Marsa (ADCM) est une association qui "
        "promeut des actions bénévoles dans la commune de La Marsa ainsi que dans d'autres "
        "localités. Elle a collaboré avec des organismes internationaux tels que l'OIM, "
        "Médecins du Monde, Décathlon et l'Alliance des Créateurs Arabes, ainsi qu'avec des "
        "intervenants locaux (SOS Gammarth, les Scouts Tunisiens, le Croissant-Rouge Tunisien), "
        "à l'occasion de nombreuses actions : Peinture pour tous, actions sociales (Hiver au "
        "chaud, nettoyage de plages, Préservons ensemble la planète, Tous contre le plastique, "
        "soutien aux sans-abri, couffins de Ramadan, visites aux maisons de retraite et aux "
        "centres pour autistes), sport pour l'inclusion (Summer Games, IntégraSport, Stay United "
        "Around Sport…), démontrant ainsi sa capacité d'organisation et sa portée associative.",
        space_after=8)

    add_paragraph(doc,
        "À travers des camps et des activités de volontariat culturel, environnemental, sportif "
        "et social, l'Association cherche à inclure les jeunes au sein de la communauté tunisienne "
        "en développant leur sens civique.",
        space_after=8)

    add_paragraph(doc, "Ces activités poursuivent les objectifs principaux suivants :",
                  space_after=4)

    add_bullet(doc,
        "l'inclusion des jeunes dans des actions de volontariat social ;")
    add_bullet(doc,
        "l'élaboration d'actions socioculturelles qui favorisent le respect de l'environnement, "
        "la non-discrimination, la responsabilité civique et la diversité culturelle ;")
    add_bullet(doc,
        "le respect de l'environnement et la promotion des valeurs écologiques lors de la mise "
        "en œuvre de toutes les activités.")

    add_paragraph(doc,
        "Ces actions s'inscrivent notamment dans le cadre de l'Objectif de Développement "
        "Durable n° 16 des Nations Unies.",
        space_before=8, space_after=10)

    add_heading(doc, "Présentation de l'Institut IPEST", level_size=12,
                space_before=8, space_after=6)

    add_paragraph(doc,
        "L'Institut Préparatoire aux Études Scientifiques et Techniques (IPEST), établissement "
        "public d'enseignement supérieur relevant de l'Université de Carthage et situé à La Marsa, "
        "assure la préparation de ses étudiants aux concours nationaux et internationaux d'accès "
        "aux écoles d'ingénieurs. Attaché à la formation de citoyens responsables autant qu'à "
        "l'excellence académique, l'Institut encourage l'engagement associatif, culturel et "
        "citoyen de ses étudiants à travers la vie de ses clubs.",
        space_after=10)

    add_paragraph(doc,
        "C'est dans ce cadre que les Parties ont décidé de collaborer, et il est convenu ce "
        "qui suit :",
        italic=True, space_after=10)

    add_separator(doc)
    add_page_break(doc)

    # ===========================================================
    # ARTICLES
    # ===========================================================

    # --- Article 1 ---
    add_heading(doc, "Article 1 — Objet de la présente convention",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc,
        "La présente convention a pour objet de définir les modalités et les conditions de la "
        "contribution des deux Parties soussignées dans le cadre d'un partenariat portant sur la "
        "création, au sein de l'Institut, du club dénommé « ADCM–IPEST », parrainé par "
        "l'Association du Développement Culturel à La Marsa.",
        space_after=12)

    # --- Article 2 ---
    add_heading(doc, "Article 2 — Modalités du partenariat",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc, "L'Association ADCM s'engage à :", bold=False, space_after=4)

    add_bullet(doc,
        "assurer des formations en soft skills, hard skills et développement personnel au profit "
        "du comité du club, et fournir les supports pédagogiques nécessaires à la mise en place "
        "de ses actions de sensibilisation ;")
    add_bullet(doc,
        "accompagner et encadrer les membres du club dans leurs actions, et les mobiliser dans "
        "les activités portées par l'Association en tant que jeunes de l'Association ADCM ;")
    add_bullet(doc,
        "participer à des actions communes au sein de l'établissement ou lors des actions "
        "culturelles, humanitaires, environnementales et sportives portées par l'Association ;")
    add_bullet(doc,
        "mentionner le nom et le logo du club dans ses supports et canaux de communication, et "
        "communiquer sur les actions communes via ces mêmes canaux.")

    add_paragraph(doc, "L'Institut IPEST s'engage à :", bold=False,
                  space_before=8, space_after=4)

    add_bullet(doc,
        "faciliter la création et le fonctionnement du club au sein de l'Institut, conformément "
        "à son règlement intérieur ;")
    add_bullet(doc,
        "désigner un représentant chargé du suivi du partenariat et de l'accompagnement du club ;")
    add_bullet(doc,
        "mettre à la disposition du club, selon ses disponibilités, des espaces pour la tenue "
        "de ses réunions et activités.")

    add_paragraph(doc, space_after=8)

    # --- Article 3 ---
    add_heading(doc, "Article 3 — Droits et prérogatives de l'Institut",
                level_size=13, space_before=8, space_after=6)

    add_bullet(doc,
        "Toute activité organisée au sein de l'Institut ou impliquant ses étudiants est soumise "
        "à l'accord préalable de l'administration de l'Institut.")
    add_bullet(doc,
        "L'utilisation du nom ou du logo de l'IPEST sur tout support de communication est "
        "subordonnée à la validation préalable de l'administration de l'Institut.")
    add_bullet(doc,
        "La participation des étudiants aux activités du club ne doit en aucun cas porter atteinte "
        "au bon déroulement de leurs études ni au fonctionnement normal de l'Institut.")
    add_bullet(doc,
        "La présente convention ne met aucune obligation financière à la charge de l'Institut.")
    add_bullet(doc,
        "L'Institut se réserve le droit de suspendre toute activité qu'il jugerait incompatible "
        "avec son règlement intérieur ou avec sa mission éducative.")

    add_paragraph(doc, space_after=8)

    # --- Article 4 ---
    add_heading(doc, "Article 4 — Prise d'effet et durée",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc,
        "La présente convention est conclue pour une durée de deux (2) ans à compter de sa date "
        "de signature. Elle est tacitement reconduite aux mêmes conditions, sauf notification "
        "préalable de l'une des Parties dans les conditions prévues à l'Article 7 ci-après.",
        space_after=12)

    # --- Article 5 ---
    add_heading(doc, "Article 5 — Renouvellement",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc,
        "La présente convention pourra faire l'objet d'un renouvellement dans les conditions "
        "définies par les deux Parties lors d'une réunion de bilan, fixée à la demande de l'une "
        "ou l'autre des Parties, permettant de faire le point sur les actions passées et les "
        "projets à venir.",
        space_after=6)

    add_paragraph(doc,
        "Le renouvellement fera alors l'objet d'un avenant spécifique précisant uniquement "
        "ces modalités.",
        space_after=6)

    add_paragraph(doc,
        "Si aucune des deux Parties ne demande de réunion de bilan avant la date d'expiration "
        "de la convention, cette dernière sera renouvelée selon les mêmes modalités.",
        space_after=12)

    # --- Article 6 ---
    add_heading(doc, "Article 6 — Adhésion au club",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc,
        "L'adhésion au club est gratuite et ouverte aux étudiants de l'Institut, sur la base "
        "du volontariat.",
        space_after=6)

    add_paragraph(doc,
        "Lorsque les actions se déroulent en dehors de l'Institut, les membres du club sont "
        "placés sous la responsabilité de l'Association ADCM, en présence d'un représentant "
        "de l'Institut.",
        space_after=12)

    # --- Article 7 ---
    add_heading(doc, "Article 7 — Résiliation",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc,
        "La présente convention peut être résiliée, pour motif légitime, par l'une ou l'autre "
        "des Parties, au moyen d'une lettre recommandée avec accusé de réception ou d'un "
        "courrier électronique précisant la cause de la résiliation, adressé au moins trente (30) "
        "jours avant la date de reconduction tacite.",
        space_after=6)

    add_paragraph(doc,
        "À tout autre moment, en cas de désaccord entre les Parties et après constat de "
        "l'impossibilité de poursuivre l'exécution dudit accord, la convention pourra être "
        "résiliée de plein droit par l'une ou l'autre des Parties, à l'expiration d'un délai "
        "de trente (30) jours suivant l'envoi d'une lettre recommandée avec accusé de réception "
        "valant mise en demeure.",
        space_after=12)

    # --- Article 8 ---
    add_heading(doc, "Article 8 — Dispositions finales",
                level_size=13, space_before=8, space_after=6)

    add_paragraph(doc,
        "Les Parties s'engagent à respecter scrupuleusement les stipulations de la présente "
        "convention. Toute modification doit être approuvée par les deux Parties et fera l'objet "
        "d'un avenant dûment signé. Tout différend relatif à l'interprétation ou à l'exécution "
        "de la présente convention sera réglé prioritairement à l'amiable entre les Parties.",
        space_after=14)

    add_separator(doc)
    add_page_break(doc)

    # ===========================================================
    # SIGNATURE PAGE
    # ===========================================================
    add_paragraph(doc, space_after=4)

    add_paragraph(doc,
        "Fait à Tunis, le ………………………………………………………",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_paragraph(doc,
        "En deux (2) exemplaires originaux, dont un remis à chacune des Parties.",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # Signature table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Remove all borders for a clean look
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    left_cell  = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    def fill_sig_cell(cell, title_line1, title_line2, name_line):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for para in cell.paragraphs:
            cell._tc.remove(para._p)

        def cp(text, bold=False, center=True, space_after=4):
            p = cell.add_paragraph()
            align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            para_format(p, alignment=align, space_before=0, space_after=space_after)
            run = p.add_run(text)
            set_font(run, bold=bold, size=11)

        cp(title_line1, bold=True)
        cp(title_line2, bold=True, space_after=6)
        cp(name_line, bold=False, space_after=8)
        cp("Lu et approuvé", bold=False, space_after=24)
        cp("Signature et cachet officiel", bold=False, space_after=0)

    fill_sig_cell(left_cell,
                  "Pour l'Institut IPEST",
                  "Monsieur/Madame ………………………………",
                  "Directeur(trice) de l'Institut")

    fill_sig_cell(right_cell,
                  "Pour l'Association ADCM",
                  "Madame Soumaya KHAMMAR",
                  "Présidente de l'Association")

    doc.save(OUTPUT_DOCX)
    print(f"DOCX saved: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


# ---------------------------------------------------------------------------
# PDF via ReportLab
# ---------------------------------------------------------------------------

def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    PageBreak, Table, TableStyle, HRFlowable)
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    # Use a bundled serif font (Helvetica as fallback since Times built-in is enough)
    doc = SimpleDocTemplate(
        OUTPUT_PDF,
        pagesize=A4,
        leftMargin=3*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        title="Convention de Partenariat ADCM–IPEST",
        author="ADCM / IPEST",
    )

    styles = getSampleStyleSheet()
    W = A4[0] - 5.5*cm  # usable width

    BODY = ParagraphStyle("body", fontName="Times-Roman", fontSize=12,
                          leading=14.5, alignment=TA_JUSTIFY,
                          spaceAfter=6, spaceBefore=0)
    BODY_ITALIC = ParagraphStyle("body_italic", parent=BODY,
                                 fontName="Times-Italic")
    HEADING = ParagraphStyle("heading", fontName="Times-Bold", fontSize=14,
                             leading=17, alignment=TA_LEFT,
                             spaceBefore=12, spaceAfter=8)
    SUBHEADING = ParagraphStyle("subheading", fontName="Times-Bold", fontSize=12,
                                leading=15, alignment=TA_LEFT,
                                spaceBefore=8, spaceAfter=6)
    ARTICLE = ParagraphStyle("article", fontName="Times-Bold", fontSize=13,
                              leading=16, alignment=TA_LEFT,
                              spaceBefore=10, spaceAfter=6)
    CENTER = ParagraphStyle("center", fontName="Times-Roman", fontSize=12,
                            leading=15, alignment=TA_CENTER,
                            spaceAfter=6)
    CENTER_BOLD = ParagraphStyle("center_bold", fontName="Times-Bold",
                                 fontSize=13, leading=16, alignment=TA_CENTER,
                                 spaceAfter=8)
    TITLE_MAIN = ParagraphStyle("title_main", fontName="Times-Bold",
                                fontSize=20, leading=24, alignment=TA_CENTER,
                                spaceAfter=12)
    BULLET = ParagraphStyle("bullet", fontName="Times-Roman", fontSize=12,
                            leading=14.5, alignment=TA_JUSTIFY,
                            leftIndent=1*cm, firstLineIndent=-0.5*cm,
                            spaceAfter=4)
    SIG = ParagraphStyle("sig", fontName="Times-Roman", fontSize=11,
                         leading=14, alignment=TA_CENTER, spaceAfter=4)
    SIG_BOLD = ParagraphStyle("sig_bold", fontName="Times-Bold", fontSize=11,
                              leading=14, alignment=TA_CENTER, spaceAfter=6)

    HR = lambda: HRFlowable(width="100%", thickness=0.6, color=colors.grey,
                             spaceAfter=8, spaceBefore=4)
    SP = lambda h=6: Spacer(1, h)

    B = lambda t: f"<b>{t}</b>"
    I = lambda t: f"<i>{t}</i>"

    story = []

    # ------ TITLE PAGE ------
    story += [SP(60),
              Paragraph("CONVENTION DE PARTENARIAT", TITLE_MAIN),
              HR(),
              SP(16),
              Paragraph("entre", CENTER),
              SP(10),
              Paragraph(B("L'Association du Développement Culturel à La Marsa<br/>« ADCM »"),
                        CENTER_BOLD),
              SP(10),
              Paragraph("et", CENTER),
              SP(10),
              Paragraph(B("L'Institut Préparatoire aux Études Scientifiques et Techniques<br/>« IPEST »"),
                        CENTER_BOLD),
              SP(30),
              HR(),
              SP(10),
              Paragraph("La Marsa — 2026", CENTER),
              PageBreak()]

    # ------ ENTRE LES SOUSSIGNÉS ------
    story += [Paragraph("ENTRE LES SOUSSIGNÉS", HEADING),
              Paragraph(
                  "L'Association du Développement Culturel à La Marsa « ADCM », représentée par "
                  "Madame Soumaya KHAMMAR, Présidente de l'Association, ayant tous pouvoirs aux "
                  "fins de signature des présentes,", BODY),
              Paragraph(I("ci-après désignée « l'Association » ou « ADCM »,"), BODY_ITALIC),
              Paragraph("d'une part,", BODY),
              SP(10),
              Paragraph("et", CENTER),
              SP(10),
              Paragraph(
                  "L'Institut Préparatoire aux Études Scientifiques et Techniques « IPEST », "
                  "représenté par Monsieur/Madame ………………………………………………………, "
                  "Directeur(trice) de l'Institut, ayant tous pouvoirs aux fins de signature "
                  "des présentes,", BODY),
              Paragraph(I("ci-après désigné « l'Institut » ou « IPEST »,"), BODY_ITALIC),
              Paragraph("d'autre part,", BODY),
              SP(10),
              Paragraph(I("Les soussignés étant ci-après désignés conjointement « les Parties »."),
                        BODY_ITALIC),
              SP(8),
              HR()]

    # ------ PRÉAMBULE ------
    story += [Paragraph("PRÉAMBULE", HEADING),
              Paragraph("Présentation de l'Association ADCM", SUBHEADING),
              Paragraph(
                  "L'Association du Développement Culturel à La Marsa (ADCM) est une association "
                  "qui promeut des actions bénévoles dans la commune de La Marsa ainsi que dans "
                  "d'autres localités. Elle a collaboré avec des organismes internationaux tels que "
                  "l'OIM, Médecins du Monde, Décathlon et l'Alliance des Créateurs Arabes, ainsi "
                  "qu'avec des intervenants locaux (SOS Gammarth, les Scouts Tunisiens, le "
                  "Croissant-Rouge Tunisien), à l'occasion de nombreuses actions : Peinture pour "
                  "tous, actions sociales (Hiver au chaud, nettoyage de plages, Préservons ensemble "
                  "la planète, Tous contre le plastique, soutien aux sans-abri, couffins de Ramadan, "
                  "visites aux maisons de retraite et aux centres pour autistes), sport pour "
                  "l'inclusion (Summer Games, IntégraSport, Stay United Around Sport…), démontrant "
                  "ainsi sa capacité d'organisation et sa portée associative.", BODY),
              Paragraph(
                  "À travers des camps et des activités de volontariat culturel, environnemental, "
                  "sportif et social, l'Association cherche à inclure les jeunes au sein de la "
                  "communauté tunisienne en développant leur sens civique.", BODY),
              Paragraph("Ces activités poursuivent les objectifs principaux suivants :", BODY),
              Paragraph("— l'inclusion des jeunes dans des actions de volontariat social ;", BULLET),
              Paragraph(
                  "— l'élaboration d'actions socioculturelles qui favorisent le respect de "
                  "l'environnement, la non-discrimination, la responsabilité civique et la "
                  "diversité culturelle ;", BULLET),
              Paragraph(
                  "— le respect de l'environnement et la promotion des valeurs écologiques lors "
                  "de la mise en œuvre de toutes les activités.", BULLET),
              Paragraph(
                  "Ces actions s'inscrivent notamment dans le cadre de l'Objectif de Développement "
                  "Durable n° 16 des Nations Unies.", BODY),
              SP(8),
              Paragraph("Présentation de l'Institut IPEST", SUBHEADING),
              Paragraph(
                  "L'Institut Préparatoire aux Études Scientifiques et Techniques (IPEST), "
                  "établissement public d'enseignement supérieur relevant de l'Université de "
                  "Carthage et situé à La Marsa, assure la préparation de ses étudiants aux "
                  "concours nationaux et internationaux d'accès aux écoles d'ingénieurs. Attaché "
                  "à la formation de citoyens responsables autant qu'à l'excellence académique, "
                  "l'Institut encourage l'engagement associatif, culturel et citoyen de ses "
                  "étudiants à travers la vie de ses clubs.", BODY),
              SP(6),
              Paragraph(
                  I("C'est dans ce cadre que les Parties ont décidé de collaborer, et il est "
                    "convenu ce qui suit :"), BODY_ITALIC),
              SP(6),
              HR(),
              PageBreak()]

    # ------ ARTICLES ------
    def art(n, title):
        return Paragraph(f"Article {n} — {title}", ARTICLE)

    story += [art(1, "Objet de la présente convention"),
              Paragraph(
                  "La présente convention a pour objet de définir les modalités et les conditions "
                  "de la contribution des deux Parties soussignées dans le cadre d'un partenariat "
                  "portant sur la création, au sein de l'Institut, du club dénommé « ADCM–IPEST », "
                  "parrainé par l'Association du Développement Culturel à La Marsa.", BODY),
              SP(10)]

    story += [art(2, "Modalités du partenariat"),
              Paragraph("L'Association ADCM s'engage à :", BODY),
              Paragraph(
                  "— assurer des formations en soft skills, hard skills et développement personnel "
                  "au profit du comité du club, et fournir les supports pédagogiques nécessaires "
                  "à la mise en place de ses actions de sensibilisation ;", BULLET),
              Paragraph(
                  "— accompagner et encadrer les membres du club dans leurs actions, et les "
                  "mobiliser dans les activités portées par l'Association en tant que jeunes de "
                  "l'Association ADCM ;", BULLET),
              Paragraph(
                  "— participer à des actions communes au sein de l'établissement ou lors des "
                  "actions culturelles, humanitaires, environnementales et sportives portées "
                  "par l'Association ;", BULLET),
              Paragraph(
                  "— mentionner le nom et le logo du club dans ses supports et canaux de "
                  "communication, et communiquer sur les actions communes via ces mêmes canaux.", BULLET),
              SP(8),
              Paragraph("L'Institut IPEST s'engage à :", BODY),
              Paragraph(
                  "— faciliter la création et le fonctionnement du club au sein de l'Institut, "
                  "conformément à son règlement intérieur ;", BULLET),
              Paragraph(
                  "— désigner un représentant chargé du suivi du partenariat et de "
                  "l'accompagnement du club ;", BULLET),
              Paragraph(
                  "— mettre à la disposition du club, selon ses disponibilités, des espaces pour "
                  "la tenue de ses réunions et activités.", BULLET),
              SP(10)]

    story += [art(3, "Droits et prérogatives de l'Institut"),
              Paragraph(
                  "— Toute activité organisée au sein de l'Institut ou impliquant ses étudiants "
                  "est soumise à l'accord préalable de l'administration de l'Institut.", BULLET),
              Paragraph(
                  "— L'utilisation du nom ou du logo de l'IPEST sur tout support de communication "
                  "est subordonnée à la validation préalable de l'administration de l'Institut.", BULLET),
              Paragraph(
                  "— La participation des étudiants aux activités du club ne doit en aucun cas "
                  "porter atteinte au bon déroulement de leurs études ni au fonctionnement normal "
                  "de l'Institut.", BULLET),
              Paragraph(
                  "— La présente convention ne met aucune obligation financière à la charge "
                  "de l'Institut.", BULLET),
              Paragraph(
                  "— L'Institut se réserve le droit de suspendre toute activité qu'il jugerait "
                  "incompatible avec son règlement intérieur ou avec sa mission éducative.", BULLET),
              SP(10)]

    story += [art(4, "Prise d'effet et durée"),
              Paragraph(
                  "La présente convention est conclue pour une durée de deux (2) ans à compter "
                  "de sa date de signature. Elle est tacitement reconduite aux mêmes conditions, "
                  "sauf notification préalable de l'une des Parties dans les conditions prévues "
                  "à l'Article 7 ci-après.", BODY),
              SP(10)]

    story += [art(5, "Renouvellement"),
              Paragraph(
                  "La présente convention pourra faire l'objet d'un renouvellement dans les "
                  "conditions définies par les deux Parties lors d'une réunion de bilan, fixée à "
                  "la demande de l'une ou l'autre des Parties, permettant de faire le point sur "
                  "les actions passées et les projets à venir.", BODY),
              Paragraph(
                  "Le renouvellement fera alors l'objet d'un avenant spécifique précisant "
                  "uniquement ces modalités.", BODY),
              Paragraph(
                  "Si aucune des deux Parties ne demande de réunion de bilan avant la date "
                  "d'expiration de la convention, cette dernière sera renouvelée selon les "
                  "mêmes modalités.", BODY),
              SP(10)]

    story += [art(6, "Adhésion au club"),
              Paragraph(
                  "L'adhésion au club est gratuite et ouverte aux étudiants de l'Institut, sur "
                  "la base du volontariat.", BODY),
              Paragraph(
                  "Lorsque les actions se déroulent en dehors de l'Institut, les membres du club "
                  "sont placés sous la responsabilité de l'Association ADCM, en présence d'un "
                  "représentant de l'Institut.", BODY),
              SP(10)]

    story += [art(7, "Résiliation"),
              Paragraph(
                  "La présente convention peut être résiliée, pour motif légitime, par l'une ou "
                  "l'autre des Parties, au moyen d'une lettre recommandée avec accusé de réception "
                  "ou d'un courrier électronique précisant la cause de la résiliation, adressé au "
                  "moins trente (30) jours avant la date de reconduction tacite.", BODY),
              Paragraph(
                  "À tout autre moment, en cas de désaccord entre les Parties et après constat de "
                  "l'impossibilité de poursuivre l'exécution dudit accord, la convention pourra "
                  "être résiliée de plein droit par l'une ou l'autre des Parties, à l'expiration "
                  "d'un délai de trente (30) jours suivant l'envoi d'une lettre recommandée avec "
                  "accusé de réception valant mise en demeure.", BODY),
              SP(10)]

    story += [art(8, "Dispositions finales"),
              Paragraph(
                  "Les Parties s'engagent à respecter scrupuleusement les stipulations de la "
                  "présente convention. Toute modification doit être approuvée par les deux Parties "
                  "et fera l'objet d'un avenant dûment signé. Tout différend relatif à "
                  "l'interprétation ou à l'exécution de la présente convention sera réglé "
                  "prioritairement à l'amiable entre les Parties.", BODY),
              SP(10),
              HR(),
              PageBreak()]

    # ------ SIGNATURE PAGE ------
    story += [SP(20),
              Paragraph("Fait à Tunis, le ………………………………………………………", CENTER),
              SP(6),
              Paragraph(
                  "En deux (2) exemplaires originaux, dont un remis à chacune des Parties.", CENTER),
              SP(30)]

    sig_data = [
        [Paragraph(B("Pour l'Institut IPEST"), SIG_BOLD),
         Paragraph(B("Pour l'Association ADCM"), SIG_BOLD)],
        [Paragraph("Monsieur/Madame ………………………………", SIG),
         Paragraph("Madame Soumaya KHAMMAR", SIG)],
        [Paragraph("Directeur(trice) de l'Institut", SIG),
         Paragraph("Présidente de l'Association", SIG)],
        [Paragraph(I("Lu et approuvé"), SIG),
         Paragraph(I("Lu et approuvé"), SIG)],
        [Spacer(1, 50), Spacer(1, 50)],
        [Paragraph("Signature et cachet officiel", SIG),
         Paragraph("Signature et cachet officiel", SIG)],
    ]

    sig_table = Table(sig_data, colWidths=[W/2, W/2])
    sig_table.setStyle(TableStyle([
        ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",     (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LINEABOVE",  (0, 4), (0, 4), 0.5, colors.black),
        ("LINEABOVE",  (1, 4), (1, 4), 0.5, colors.black),
    ]))
    story.append(sig_table)

    doc.build(story)
    print(f"PDF saved: {OUTPUT_PDF}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print("Done.")
